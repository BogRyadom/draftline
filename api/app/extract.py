"""Document text extraction and chunking for the knowledge base.

Pure functions (no network, no DB) so they are easy to unit-test. Chunk size is
~800 tokens approximated as characters (~4 chars/token) with overlap, split on
whitespace/paragraph boundaries.
"""

from __future__ import annotations

import io
import re

from docx import Document as DocxDocument
from pypdf import PdfReader

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_MIME_TYPES = {PDF_MIME, DOCX_MIME}

_CHUNK_CHARS = 3200  # ~800 tokens
_OVERLAP_CHARS = 400  # ~100 tokens


class UnsupportedFileType(Exception):
    """Raised for a file that is neither PDF nor DOCX."""


def extract_text(data: bytes, *, mime_type: str | None, filename: str) -> str:
    """Extract plain text from a PDF or DOCX byte payload."""
    mt = (mime_type or "").lower()
    name = (filename or "").lower()
    if mt == PDF_MIME or name.endswith(".pdf"):
        return _extract_pdf(data)
    if mt == DOCX_MIME or name.endswith(".docx"):
        return _extract_docx(data)
    raise UnsupportedFileType(f"Unsupported file type: {mime_type or filename!r}")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def normalize_text(text: str) -> str:
    """Collapse runaway whitespace while preserving paragraph breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(
    text: str, *, chunk_chars: int = _CHUNK_CHARS, overlap_chars: int = _OVERLAP_CHARS
) -> list[str]:
    """Split text into overlapping chunks, breaking on whitespace where possible."""
    text = normalize_text(text)
    if not text:
        return []
    if len(text) <= chunk_chars:
        return [text]

    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_chars, n)
        if end < n:
            # Prefer a clean break within the last stretch of the window.
            floor = max(start + chunk_chars - 300, start + 1)
            cut = max(text.rfind(" ", floor, end), text.rfind("\n", floor, end))
            if cut > start:
                end = cut
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap_chars, start + 1)
    return chunks
